import io
import copy
from docx import Document
from app.pii_engine import sanitize


def parse_docx(data: bytes, mode: str = "redact") -> tuple[bytes, int]:
    doc = Document(io.BytesIO(data))

    # Collect all paragraphs (body + tables) with their text
    all_paras = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_paras.append(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        all_paras.append(para)

    if not all_paras:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), 0

    # Batch: join all paragraph texts, single sanitize call
    separator = "\n\x00\n"
    all_text = separator.join(p.text for p in all_paras)
    cleaned_text, total = sanitize(all_text, mode)

    if total > 0:
        cleaned_parts = cleaned_text.split(separator)
        for idx, para in enumerate(all_paras):
            if idx < len(cleaned_parts) and cleaned_parts[idx] != para.text:
                _replace_runs(para, cleaned_parts[idx])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), total


def _replace_runs(paragraph, new_text: str):
    if not paragraph.runs:
        paragraph.text = new_text
        return
    fmt = copy.deepcopy(paragraph.runs[0].font)
    for i in range(len(paragraph.runs) - 1, 0, -1):
        paragraph.runs[i].clear()
        r_element = paragraph.runs[i]._element
        r_element.getparent().remove(r_element)
    paragraph.runs[0].text = new_text
    run = paragraph.runs[0]
    if fmt.bold is not None:
        run.font.bold = fmt.bold
    if fmt.italic is not None:
        run.font.italic = fmt.italic
    if fmt.size is not None:
        run.font.size = fmt.size
    if fmt.name is not None:
        run.font.name = fmt.name
